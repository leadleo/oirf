"""
Generate report_framework.docx from report_framework_handoff.json.

Style spec (derived from reference docx):
  Font: Calibri throughout
  Page margins: 1 inch all sides
  Cover:
    - 5 blank paragraphs
    - Title: 24pt, bold, #0B2545, center
    - Subtitle: 13pt, #555555, center
    - Version / Date: 11pt, #555555, center
  TOC page:
    - "目录" Heading 1: 16pt, bold, #2E74B5
    - Chapter entries: 11pt, black, left
    - Subsection entries: 11pt, black, indented ~4 chars
  Body:
    - Heading 1 (chapter): 16pt, bold, #2E74B5, left
    - Heading 2 (subsection): 13pt, bold, #2E74B5, left
    - Description paragraph: 11pt, black, left
    - "预期输出" line: 10.5pt, bold, #1F4D78, left
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Colour constants ──────────────────────────────────────────────────────────
NAVY       = RGBColor(0x0B, 0x25, 0x45)   # cover title
GREY       = RGBColor(0x55, 0x55, 0x55)   # cover subtitle / version / date
BLUE       = RGBColor(0x2E, 0x74, 0xB5)   # headings / TOC title
DARK_BLUE  = RGBColor(0x1F, 0x4D, 0x78)   # 预期输出
BLACK      = RGBColor(0x00, 0x00, 0x00)

# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_run(run, *, size: Pt, color: RGBColor, bold: bool | None = None, font_name: str = "Calibri"):
    run.font.size = size
    run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_name


def _add_cover(doc, *, title: str, subtitle: str, version: str, date: str):
    """Append 5 blank paragraphs then the cover block."""
    for _ in range(5):
        doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run(run, size=Pt(24), color=NAVY, bold=True)

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        _set_run(run, size=Pt(13), color=GREY)

    # Version
    if version:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(version)
        _set_run(run, size=Pt(11), color=GREY)

    # Date
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date)
        _set_run(run, size=Pt(11), color=GREY)


def _add_toc_page(doc, chapters: list[dict]):
    """Manual table-of-contents page with improved styling."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Page break before TOC ──────────────────────────────────────────────────
    doc.add_page_break()

    # ── "目录" section heading ─────────────────────────────────────────────────
    h = doc.add_heading("目录", level=1)
    for run in h.runs:
        _set_run(run, size=Pt(18), color=BLUE, bold=True)

    # Decorative underline: thin blue border-bottom on the heading paragraph
    pPr = h._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")       # 1pt
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Spacer after heading
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    # ── Chapter + subsection entries ───────────────────────────────────────────
    for chapter in chapters:
        # Chapter line: bold, navy colour, extra top spacing to separate from previous block
        p_ch = doc.add_paragraph()
        p_ch.paragraph_format.space_before = Pt(6)
        p_ch.paragraph_format.space_after  = Pt(2)
        run = p_ch.add_run(chapter.get("chapter_title", ""))
        _set_run(run, size=Pt(11), color=NAVY, bold=True)

        # Subsection lines: indented, grey colour, no bold
        for sub in chapter.get("external_subsections", []):
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.left_indent  = Cm(1.0)
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after  = Pt(0)
            label = sub if isinstance(sub, str) else sub.get("title", str(sub))
            run = p_sub.add_run(label)
            _set_run(run, size=Pt(10.5), color=GREY, bold=False)


def _add_body(doc, chapters: list[dict]):
    """Main report-framework body. Starts on a new page."""
    doc.add_page_break()

    for chapter in chapters:
        # ── Chapter heading ────────────────────────────────────────────────────
        h1 = doc.add_heading(chapter.get("chapter_title", ""), level=1)
        for run in h1.runs:
            _set_run(run, size=Pt(16), color=BLUE, bold=True)

        # ── Subsections ────────────────────────────────────────────────────────
        subsections = chapter.get("external_subsections", [])
        descriptions = chapter.get("descriptions", {})   # optional map: subsection_title → description
        placeholders = chapter.get("output_placeholders", [])

        for idx, sub in enumerate(subsections):
            label = sub if isinstance(sub, str) else sub.get("title", str(sub))
            h2 = doc.add_heading(label, level=2)
            for run in h2.runs:
                _set_run(run, size=Pt(13), color=BLUE, bold=True)

            # Description if provided (matched by index or by title key)
            desc = None
            if isinstance(descriptions, dict):
                desc = descriptions.get(label)
            elif isinstance(descriptions, list) and idx < len(descriptions):
                desc = descriptions[idx]
            if desc:
                p = doc.add_paragraph()
                run = p.add_run(desc)
                _set_run(run, size=Pt(11), color=BLACK)

        # ── Expected-output lines ──────────────────────────────────────────────
        for ph in placeholders:
            text = ph if isinstance(ph, str) else ph.get("text", str(ph))
            if not text:
                continue
            # Prefix with "预期输出：" if not already present
            display = text if text.startswith("预期输出") else f"预期输出：{text}"
            p = doc.add_paragraph()
            run = p.add_run(display)
            _set_run(run, size=Pt(10.5), color=DARK_BLUE, bold=True)


# ── Main ──────────────────────────────────────────────────────────────────────

def build_docx(data: dict, output_path: Path):
    doc = Document()

    # ── Page margins (1 inch all sides) ────────────────────────────────────────
    for section in doc.sections:
        section.left_margin   = Emu(914400)
        section.right_margin  = Emu(914400)
        section.top_margin    = Emu(914400)
        section.bottom_margin = Emu(914400)

    # ── Extract data ───────────────────────────────────────────────────────────
    framework   = data.get("report_framework", {})
    title       = framework.get("external_title", "报告框架")
    subtitle    = framework.get("subtitle", "")
    version     = data.get("version", "框架稿")
    date_str    = data.get("date", "")
    chapters    = framework.get("chapter_outline", [])

    # ── Build ──────────────────────────────────────────────────────────────────
    _add_cover(doc, title=title, subtitle=subtitle, version=version, date=date_str)
    _add_toc_page(doc, chapters)
    _add_body(doc, chapters)

    doc.save(str(output_path))
    print(f"[OK] saved → {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_report_docx.py <input.json> [output.docx]")
        sys.exit(1)

    input_path  = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else input_path.with_name("report_framework.docx")

    data = json.loads(input_path.read_text(encoding="utf-8"))
    build_docx(data, output_path)


if __name__ == "__main__":
    main()
